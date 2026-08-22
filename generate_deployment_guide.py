#!/usr/bin/env python3
"""
Generate comprehensive SherryJo Calendar App Deployment & Troubleshooting Guide (DOCX)
A-Team Production-Grade Documentation
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ============================================================================
# TITLE & METADATA
# ============================================================================
title = doc.add_paragraph()
title_run = title.add_run('SherryJo Calendar App\nFull Stack Deployment & Troubleshooting Guide')
title_run.font.size = Pt(24)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

meta = doc.add_paragraph()
meta_text = f'Last Updated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\nProduction Environment | A-Team Grade Documentation'
meta_run = meta.add_run(meta_text)
meta_run.font.size = Pt(10)
meta_run.font.italic = True
meta_run.font.color.rgb = RGBColor(100, 100, 100)
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_page_break()

# ============================================================================
# TABLE OF CONTENTS
# ============================================================================
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. Architecture Overview',
    '3. Platform Status & Health Checks',
    '4. What Just Happened (Recent Issues)',
    '5. Quick Fix Checklist',
    '6. Detailed Troubleshooting Guide',
    '7. Full Redeployment Procedures',
    '8. Git Workflow & Commit/Push',
    '9. Emergency Recovery Procedures',
    '10. Monitoring & Alerts',
    '11. Command Reference',
]
for item in toc_items:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)

doc.add_page_break()

# ============================================================================
# 1. EXECUTIVE SUMMARY
# ============================================================================
doc.add_heading('1. Executive Summary', level=1)

doc.add_paragraph(
    'The SherryJo Calendar application runs across 4 interconnected platforms. When one '
    'component breaks, the entire user experience fails. This guide provides step-by-step '
    'procedures to diagnose, fix, and redeploy all components.'
)

doc.add_heading('Status Dashboard', level=2)
table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'
header_cells = table.rows[0].cells
header_cells[0].text = 'Platform'
header_cells[1].text = 'Status Check URL'
header_cells[2].text = 'Purpose'

rows_data = [
    ('Cloudflare Workers', 'https://sherryjo-cal-app.realty-cal.workers.dev/login', 'Frontend/UI delivery'),
    ('Render Backend', 'https://sherryjo-cal-app.onrender.com/health', 'API & database'),
    ('GitHub', 'github.com/cjdawgs/SherryJo_Cal_App', 'Source control'),
    ('Supabase', 'supabase.com console', 'Database & auth'),
]

for i, (platform, url, purpose) in enumerate(rows_data, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = platform
    row_cells[1].text = url
    row_cells[2].text = purpose

doc.add_page_break()

# ============================================================================
# 2. ARCHITECTURE OVERVIEW
# ============================================================================
doc.add_heading('2. Architecture Overview', level=1)

doc.add_heading('System Flow', level=2)
flow_para = doc.add_paragraph()
flow_para.add_run('User Browser → Cloudflare Workers → Render Backend → Supabase').bold = True

doc.add_paragraph(
    'Each layer has specific responsibilities:\n\n'
    '• Cloudflare Workers: Serves HTML/CSS/JS files, routes requests\n'
    '• Render Backend: FastAPI server, handles authentication, API calls\n'
    '• Supabase: PostgreSQL database, manages calendar data\n'
    '• GitHub: Source code repository, triggers CI/CD'
)

doc.add_heading('Key Files & Locations', level=2)
locations = [
    ('wrangler.toml', 'Cloudflare Worker configuration (asset serving, secrets)'),
    ('platform/cloudflare/src/worker.js', 'Main Worker logic (1500+ lines)'),
    ('platform/cloudflare/.worker-assets/', 'Static files: HTML, CSS, JS'),
    ('app/main.py', 'Render FastAPI backend entry point'),
    ('app/routers/calendar.py', 'Calendar API endpoints'),
    ('requirements.txt', 'Python dependencies'),
]

for filename, description in locations:
    p = doc.add_paragraph(f'{filename}: {description}', style='List Bullet')

doc.add_page_break()

# ============================================================================
# 3. PLATFORM STATUS & HEALTH CHECKS
# ============================================================================
doc.add_heading('3. Platform Status & Health Checks', level=1)

doc.add_heading('Cloudflare Workers Health Check', level=2)
doc.add_paragraph('Run these commands to verify Cloudflare deployment:')
code_block = doc.add_paragraph()
code_run = code_block.add_run(
    'curl -I https://sherryjo-cal-app.realty-cal.workers.dev/login\n'
    'Expected: HTTP/2 200 with content-type: text/html\n\n'
    'curl -I https://sherryjo-cal-app.realty-cal.workers.dev/static/calendar.js\n'
    'Expected: HTTP/2 200 with content-type: text/javascript'
)
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)

doc.add_heading('Render Backend Health Check', level=2)
code_block = doc.add_paragraph()
code_run = code_block.add_run(
    'curl https://sherryjo-cal-app.onrender.com/health\n'
    'Expected: {"status":"ok","app":"running","schema_status":"ok"}'
)
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)

doc.add_heading('Database Connection Test', level=2)
doc.add_paragraph('From Render dashboard:\n1. Go to Dashboard → sherryjo-cal-app\n2. Check Deploy Log for any database connection errors\n3. Look for "schema_status": "ok" in health check response')

doc.add_page_break()

# ============================================================================
# 4. WHAT JUST HAPPENED (RECENT ISSUES)
# ============================================================================
doc.add_heading('4. What Just Happened (Recent Issues)', level=1)

doc.add_heading('The 307 Redirect Loop Issue', level=2)
doc.add_paragraph(
    'Symptoms:\n'
    '• User clicks login → infinite redirect loop\n'
    '• Browser shows: NS_ERROR_REDIRECT_LOOP\n'
    '• /login returns HTTP 307 location: /login\n\n'
    'Root Cause:\n'
    'The wrangler.toml [assets] section had TWO conflicting settings:\n'
    '1. html_handling = "none" (disabled HTML handling)\n'
    '2. run_worker_first = true (bypassed asset serving)\n\n'
    'These settings forced Cloudflare to use env.ASSETS.fetch() which was misconfigured, '
    'causing HTML pages to return 307 redirects instead of the actual content.'
)

doc.add_heading('Partial Asset Serving', level=2)
doc.add_paragraph(
    'Interesting observation:\n'
    '• /static/calendar.js worked (returned HTTP 200)\n'
    '• /login.html failed (returned HTTP 503 "binding not configured")\n'
    '• /login.html redirect to self\n\n'
    'This indicated the asset binding was partially broken, affecting only HTML routes '
    'mapped in NATIVE_PAGE_ASSETS.'
)

doc.add_page_break()

# ============================================================================
# 5. QUICK FIX CHECKLIST
# ============================================================================
doc.add_heading('5. Quick Fix Checklist', level=1)

doc.add_paragraph('If the app is down right now, follow this priority order:')

doc.add_heading('Step 1: Verify Current Status (30 seconds)', level=2)
steps = [
    'Run: curl -I https://sherryjo-cal-app.realty-cal.workers.dev/login',
    'If HTTP 200 → App is working, skip to Step 4',
    'If HTTP 307 or 503 → Continue to Step 2',
]
for step in steps:
    doc.add_paragraph(step, style='List Number')

doc.add_heading('Step 2: Check wrangler.toml (1 minute)', level=2)
doc.add_paragraph(
    'Edit /workspaces/SherryJo_Cal_App/wrangler.toml\n\n'
    'Look for the [assets] section. It should be:\n'
)
code_block = doc.add_paragraph()
code_run = code_block.add_run(
    '[assets]\n'
    'directory = "platform/cloudflare/.worker-assets"\n'
    'binding = "ASSETS"'
)
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)

doc.add_paragraph(
    '\nIf it contains:\n'
    '• html_handling = "none" → DELETE THIS LINE\n'
    '• run_worker_first = true → DELETE THIS LINE\n'
    '• Any other settings in [assets] → Ask what they do'
)

doc.add_heading('Step 3: Redeploy (2 minutes)', level=2)
code_block = doc.add_paragraph()
code_run = code_block.add_run(
    'cd /workspaces/SherryJo_Cal_App\n'
    'git add wrangler.toml\n'
    'git commit -m "Fix: Remove conflicting asset settings"\n'
    'git push\n'
    'wrangler deploy'
)
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)

doc.add_paragraph('Then run the health check again (Step 1).')

doc.add_heading('Step 4: If App Is Working', level=2)
doc.add_paragraph('Great! Continue to Section 6 to diagnose what broke it and prevent future incidents.')

doc.add_page_break()

# ============================================================================
# 6. DETAILED TROUBLESHOOTING GUIDE
# ============================================================================
doc.add_heading('6. Detailed Troubleshooting Guide', level=1)

issues = [
    {
        'title': 'Login Page Shows Blank / Infinite Redirect',
        'symptoms': [
            '/login returns HTTP 307 redirecting to /login',
            'NS_ERROR_REDIRECT_LOOP in browser console',
            'curl -I /login shows HTTP 307, location: /login',
        ],
        'diagnosis': [
            'Check wrangler.toml [assets] section for "run_worker_first = true"',
            'Check for "html_handling = none" setting',
            'Verify platform/cloudflare/.worker-assets/login.html exists (ls -la)',
        ],
        'fix': [
            'Remove offending lines from wrangler.toml',
            'Run: wrangler deploy',
            'Wait 10 seconds for edge cache to clear',
            'Test: curl -I https://sherryjo-cal-app.realty-cal.workers.dev/login',
        ],
    },
    {
        'title': 'Static Assets Not Loading (CSS/JS Blank)',
        'symptoms': [
            'Page loads but no styling or functionality',
            'Browser console shows 404 on /static/* files',
            'calendar.js fails to load',
        ],
        'diagnosis': [
            'Run: curl -I https://sherryjo-cal-app.realty-cal.workers.dev/static/calendar.js',
            'If HTTP 404: Asset files not deployed',
            'If HTTP 503: Asset binding misconfigured',
            'Check: ls -la platform/cloudflare/.worker-assets/static/',
        ],
        'fix': [
            'Verify files exist locally: ls -la platform/cloudflare/.worker-assets/static/',
            'If missing: git restore platform/cloudflare/.worker-assets/',
            'Redeploy: wrangler deploy',
            'Manually test each file: curl -I /static/calendar.js /static/style.css',
        ],
    },
    {
        'title': 'API Calls Fail (Backend Not Responding)',
        'symptoms': [
            'Login fails even with correct credentials',
            'Calendar events don\'t load',
            'Network tab shows 502 or 503 on API calls',
        ],
        'diagnosis': [
            'Check Render health: curl https://sherryjo-cal-app.onrender.com/health',
            'Check Render deployment logs in dashboard',
            'Verify Supabase database is accessible',
            'Look for connection string errors in logs',
        ],
        'fix': [
            'If Render is down: Trigger redeploy from Render dashboard',
            'If database error: Check Supabase connection pool status',
            'Check environment variables in Render dashboard match .env.local',
            'Restart Render service: Dashboard → Services → Restart',
        ],
    },
    {
        'title': 'Deployed Code Doesn\'t Match GitHub',
        'symptoms': [
            'You made changes locally, they don\'t appear on production',
            'curl shows old cached version',
            'Browser cache not clearing',
        ],
        'diagnosis': [
            'Run: git status (should be clean)',
            'Run: git log -1 --oneline (should show latest commit)',
            'Verify commit is pushed: git push (should say "Everything up-to-date")',
            'Check Cloudflare version ID: wrangler deploy (shows version hash)',
        ],
        'fix': [
            'Commit & push all changes: git add -A && git commit -m "..." && git push',
            'Full redeploy: wrangler deploy',
            'Clear browser cache: Ctrl+Shift+Del → Clear recent',
            'Hard refresh: Ctrl+Shift+R (not just Ctrl+R)',
        ],
    },
    {
        'title': 'Database Queries Timing Out',
        'symptoms': [
            'Events load slowly or not at all',
            'Supabase dashboard shows high query count',
            'Render logs show timeout errors',
        ],
        'diagnosis': [
            'Check Supabase connection pool: Supabase → Settings → Database',
            'Run query directly: psql command (if you have credentials)',
            'Check for N+1 queries in app/routers/calendar.py',
            'Monitor: Supabase → Monitoring tab',
        ],
        'fix': [
            'Optimize database queries (review N+1 issues)',
            'Increase connection pool size: Supabase → Settings',
            'Add database indexes: Run alembic migration',
            'Restart Render service to clear connection pool',
        ],
    },
]

for idx, issue in enumerate(issues, 1):
    doc.add_heading(f'Issue {idx}: {issue["title"]}', level=2)
    
    doc.add_heading('Symptoms', level=3)
    for symptom in issue['symptoms']:
        doc.add_paragraph(symptom, style='List Bullet')
    
    doc.add_heading('Diagnosis Steps', level=3)
    for step in issue['diagnosis']:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_heading('Fix', level=3)
    for fix_step in issue['fix']:
        doc.add_paragraph(fix_step, style='List Number')
    
    doc.add_paragraph()

doc.add_page_break()

# ============================================================================
# 7. FULL REDEPLOYMENT PROCEDURES
# ============================================================================
doc.add_heading('7. Full Redeployment Procedures', level=1)

doc.add_heading('7.1 Redeploy Cloudflare Workers Only', level=2)

doc.add_paragraph('Use this if you\'ve made code changes to Worker logic or updated wrangler.toml:')
code_block = doc.add_paragraph()
code_run = code_block.add_run(
    'cd /workspaces/SherryJo_Cal_App\n'
    'git add -A\n'
    'git commit -m "Update: [describe changes]"\n'
    'git push\n'
    'wrangler deploy'
)
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)

doc.add_paragraph('\nExpected output: "Deployed sherryjo-cal-app triggers (X.XX sec)"')
doc.add_paragraph('Then verify: curl -I https://sherryjo-cal-app.realty-cal.workers.dev/login')

doc.add_heading('7.2 Redeploy Render Backend', level=2)

doc.add_paragraph('Use this if you\'ve changed Python code, dependencies, or database migrations:')

doc.add_heading('Option A: Automatic (Recommended)', level=3)
doc.add_paragraph(
    '1. Git add, commit, push your changes to main\n'
    '2. Go to render.com dashboard\n'
    '3. Select "sherryjo-cal-app" service\n'
    '4. Click "Deploy → Deploy latest commit"\n'
    '5. Watch Deploy Log until you see "Deploy successful"'
)

doc.add_heading('Option B: Manual Restart', level=3)
doc.add_paragraph(
    '1. Go to render.com dashboard\n'
    '2. Select "sherryjo-cal-app"\n'
    '3. Click "Restart"'
)

doc.add_paragraph('Then verify: curl https://sherryjo-cal-app.onrender.com/health')

doc.add_heading('7.3 Redeploy Everything (Nuclear Option)', level=2)

doc.add_paragraph('Use this only if you\'re not sure what broke. Resets all platforms to latest code:')
code_block = doc.add_paragraph()
code_run = code_block.add_run(
    '# 1. Verify local state is clean\n'
    'cd /workspaces/SherryJo_Cal_App\n'
    'git status  # should be clean\n'
    'git log -1 --oneline  # should show latest commit\n\n'
    '# 2. Pull latest from GitHub (in case another dev pushed)\n'
    'git pull\n\n'
    '# 3. Redeploy Cloudflare\n'
    'wrangler deploy\n'
    'sleep 10  # wait for edge cache\n\n'
    '# 4. Redeploy Render (via dashboard button or CLI)\n'
    '# Dashboard: render.com → Deploy → Deploy latest commit\n\n'
    '# 5. Verify all platforms\n'
    'curl -I https://sherryjo-cal-app.realty-cal.workers.dev/login\n'
    'curl https://sherryjo-cal-app.onrender.com/health'
)
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)

doc.add_page_break()

# ============================================================================
# 8. GIT WORKFLOW & COMMIT/PUSH
# ============================================================================
doc.add_heading('8. Git Workflow & Commit/Push', level=1)

doc.add_heading('Before You Start Any Changes', level=2)
code_block = doc.add_paragraph()
code_run = code_block.add_run(
    'cd /workspaces/SherryJo_Cal_App\n'
    'git status  # Make sure nothing uncommitted\n'
    'git pull    # Get latest from GitHub'
)
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)

doc.add_heading('Making and Committing Changes', level=2)

commit_steps = [
    'Edit the file(s) you need to change',
    'Review changes: git diff [filename]',
    'Stage changes: git add [filename] or git add -A for all',
    'Commit: git commit -m "Clear message describing what changed"',
    'Push: git push',
]
for step in commit_steps:
    doc.add_paragraph(step, style='List Number')

doc.add_heading('Good Commit Messages', level=2)
doc.add_paragraph('Follow this format: [Type]: [Description]\n\nExamples:')

examples = [
    'Fix: Remove run_worker_first setting from wrangler.toml',
    'Feature: Add calendar event color coding',
    'Refactor: Optimize database queries in calendar.py',
    'Docs: Update troubleshooting guide with new procedures',
]
for example in examples:
    doc.add_paragraph(example, style='List Bullet')

doc.add_heading('Common Git Commands', level=2)
git_commands = [
    ('git status', 'See what files have changed'),
    ('git diff', 'See exact changes before committing'),
    ('git log -10 --oneline', 'See last 10 commits'),
    ('git push', 'Send commits to GitHub'),
    ('git pull', 'Get latest changes from GitHub'),
    ('git restore [file]', 'Undo changes to a file'),
    ('git reset HEAD~1', 'Undo last commit (keep changes)'),
]

for cmd, description in git_commands:
    p = doc.add_paragraph()
    p.add_run(cmd).font.name = 'Courier New'
    p.add_run(f': {description}')

doc.add_page_break()

# ============================================================================
# 9. EMERGENCY RECOVERY PROCEDURES
# ============================================================================
doc.add_heading('9. Emergency Recovery Procedures', level=1)

doc.add_heading('App Completely Down (Nothing Loading)', level=2)

doc.add_paragraph('Step 1: Check Cloudflare Worker Status')
code_block = doc.add_paragraph()
code_run = code_block.add_run('wrangler deployments list  # See recent deployments')
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)

doc.add_paragraph('If latest deployment looks wrong:')
code_block = doc.add_paragraph()
code_run = code_block.add_run(
    'wrangler rollback  # Automatically rollback to previous working version\n'
    'sleep 10\n'
    'curl -I https://sherryjo-cal-app.realty-cal.workers.dev/login'
)
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)

doc.add_paragraph('Step 2: If Rollback Doesn\'t Work, Restore from Git')
code_block = doc.add_paragraph()
code_run = code_block.add_run(
    'cd /workspaces/SherryJo_Cal_App\n'
    'git log --oneline  # Find last known good commit\n'
    'git reset --hard [commit_hash]  # Go back to that commit\n'
    'git push --force-with-lease\n'
    'wrangler deploy'
)
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)

doc.add_heading('Database Corrupted / Supabase Down', level=2)
doc.add_paragraph(
    '1. Check Supabase status: https://status.supabase.com\n'
    '2. If status page shows issues, wait for their resolution\n'
    '3. If status looks good but your app can\'t connect:\n'
    '   a. Check DATABASE_URL in Render environment variables\n'
    '   b. Verify Supabase connection pool hasn\'t hit limits\n'
    '   c. Restart Render service\n'
    '4. If nothing works, contact Supabase support with your project URL'
)

doc.add_heading('Lost or Corrupted Code', level=2)
doc.add_paragraph(
    'You accidentally deleted important code? Don\'t panic:\n\n'
    '1. Check git status: git status\n'
    '2. Restore file: git restore [filename]\n'
    '3. Or restore entire directory: git restore .\n'
    '4. Verify: git status (should be clean)\n'
    '5. If you already committed bad code:\n'
    '   a. git log --oneline (find last good commit)\n'
    '   b. git reset --hard [good_commit_hash]\n'
    '   c. git push --force-with-lease (warns if someone else pushed)'
)

doc.add_page_break()

# ============================================================================
# 10. MONITORING & ALERTS
# ============================================================================
doc.add_heading('10. Monitoring & Alerts', level=1)

doc.add_heading('Daily Health Checks', level=2)

doc.add_paragraph('Add these to your daily routine:')

daily_checks = [
    'Cloudflare: curl -I https://sherryjo-cal-app.realty-cal.workers.dev/login',
    'Backend: curl https://sherryjo-cal-app.onrender.com/health',
    'GitHub: git log -1 --oneline (latest commit still deployed?)',
    'Render Dashboard: Check "Deploy Log" for any recent errors',
]
for check in daily_checks:
    doc.add_paragraph(check, style='List Bullet')

doc.add_heading('What to Watch For', level=2)

warnings = [
    'Increased HTTP 307 or 503 responses',
    'Deploy times suddenly increasing',
    'Render showing "Build failed" in logs',
    'Supabase connection pool exhausted warnings',
    'Cloudflare showing high error rates in analytics',
]

for warning in warnings:
    doc.add_paragraph(warning, style='List Bullet')

doc.add_heading('Monitoring URLs', level=2)

monitor_urls = [
    ('Cloudflare Analytics', 'https://dash.cloudflare.com → yourworker → Analytics'),
    ('Render Logs', 'https://dashboard.render.com → sherryjo-cal-app → Logs'),
    ('Supabase Status', 'https://supabase.com → Monitoring'),
]

for name, url in monitor_urls:
    p = doc.add_paragraph()
    p.add_run(name).bold = True
    p.add_run(f': {url}')

doc.add_page_break()

# ============================================================================
# 11. COMMAND REFERENCE
# ============================================================================
doc.add_heading('11. Command Reference', level=1)

doc.add_heading('Cloudflare Wrangler Commands', level=2)
wrangler_cmds = [
    ('wrangler deploy', 'Deploy latest code to Cloudflare Workers'),
    ('wrangler deployments list', 'See deployment history'),
    ('wrangler rollback', 'Rollback to previous deployment'),
    ('wrangler tail', 'Watch real-time logs from Worker'),
]

for cmd, desc in wrangler_cmds:
    p = doc.add_paragraph()
    p.add_run(cmd).font.name = 'Courier New'
    p.add_run(f': {desc}')

doc.add_heading('Testing & Verification Commands', level=2)
test_cmds = [
    ('curl -I https://sherryjo-cal-app.realty-cal.workers.dev/login', 'Check login page status'),
    ('curl -I https://sherryjo-cal-app.realty-cal.workers.dev/static/calendar.js', 'Check JS file'),
    ('curl https://sherryjo-cal-app.onrender.com/health', 'Check backend health'),
    ('git status', 'See current state of git'),
    ('git log -1 --oneline', 'See latest commit'),
]

for cmd, desc in test_cmds:
    p = doc.add_paragraph()
    p.add_run(cmd).font.name = 'Courier New'
    p.add_run(f'\n{desc}')

doc.add_heading('Git Commands', level=2)
git_cmds = [
    ('git add [file]', 'Stage changes'),
    ('git add -A', 'Stage all changes'),
    ('git commit -m "message"', 'Commit with message'),
    ('git push', 'Push to GitHub'),
    ('git pull', 'Pull from GitHub'),
    ('git restore [file]', 'Undo changes to file'),
    ('git log --oneline', 'See commit history'),
]

for cmd, desc in git_cmds:
    p = doc.add_paragraph()
    p.add_run(cmd).font.name = 'Courier New'
    p.add_run(f': {desc}')

doc.add_page_break()

# ============================================================================
# CONTACT & SUPPORT
# ============================================================================
doc.add_heading('Support & Escalation', level=1)

doc.add_paragraph(
    'If you\'ve followed all steps above and the app is still broken:\n\n'
    '1. Check GitHub Issues (github.com/cjdawgs/SherryJo_Cal_App/issues)\n'
    '2. Review Render logs for Python stack traces (render.com dashboard)\n'
    '3. Check Supabase status page for database issues\n'
    '4. Consult Cloudflare documentation: https://developers.cloudflare.com/workers\n'
    '5. Enable wrangler tail to watch live Worker logs: wrangler tail'
)

doc.add_paragraph()
footer = doc.add_paragraph()
footer_text = 'End of Document | Last Updated: ' + datetime.now().strftime("%Y-%m-%d")
footer_run = footer.add_run(footer_text)
footer_run.font.size = Pt(9)
footer_run.font.italic = True
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============================================================================
# SAVE DOCUMENT
# ============================================================================
output_path = '/workspaces/SherryJo_Cal_App/DEPLOYMENT_TROUBLESHOOTING_GUIDE.docx'
doc.save(output_path)
print(f'✓ Document created: {output_path}')
